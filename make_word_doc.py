"""

    """

import pandas as pd
from docx import Document
from collections import OrderedDict
from docx.shared import Pt

FP = '/Users/mmir/Library/CloudStorage/Dropbox/git/240822_authors_list_spreadsheet_meta_analysis_paper/0-Authors List-Meta Analysis paper-Aug24.xlsx'

def one_author_affs(author, affiliations_dict, author_affindex) :
    affil_indices = []

    aff_lvls = {
            'Primary Affiliation' : None,
            'Second Affiliation'  : None,
            'Third Affiliation'   : None,
            'Fourth Affiliation'  : None,
            }

    for aff_lvl in aff_lvls :
        aff = author[aff_lvl]
        if pd.notna(aff) and aff not in [' ', ''] :
            if aff not in affiliations_dict :
                affiliations_dict[aff] = len(affiliations_dict) + 1
            affil_indices.append(affiliations_dict[aff])

    name = 'Name W/O Title (Name Actually Used in Paper Author List)'
    affil_indices.sort()
    author_affindex.append((author[name], affil_indices))

    return affiliations_dict, author_affindex

def main() :
    pass

    ##
    df = pd.read_excel(FP)

    ##
    authors = df.to_dict(orient = 'records')

    print(authors)

    ##

    # Dictionary to store unique affiliations
    affiliations_dict = OrderedDict()

    # List to store formatted author entries
    author_affindex = []

    # Process each author
    for author in authors :
        affiliations_dict, author_affindex = one_author_affs(author,
                                                             affiliations_dict,
                                                             author_affindex)

    ##
    print(affiliations_dict)

    ##
    print(author_affindex)

    ##
    # Create a new Word document
    doc = Document()

    # Create the authors list paragraph
    authors_paragraph = doc.add_paragraph()
    authors_paragraph.style.font.size = Pt(9)

    for author_name, aff_indices in author_affindex :
        run = authors_paragraph.add_run(f"{author_name}")
        if aff_indices :
            for i, idx in enumerate(aff_indices) :
                if i > 0 :
                    run = authors_paragraph.add_run(",")
                    run.font.superscript = True  # Set the superscript
                run = authors_paragraph.add_run(f"{idx}")
                run.font.superscript = True  # Set the superscript

        if author_name != author_affindex[-1][0] :
            authors_paragraph.add_run(", ")

    affliations_paragraph = doc.add_paragraph()
    affliations_paragraph.style.font.size = Pt(9)

    # Create the affiliations list
    for affil, index in affiliations_dict.items() :
        affliations_paragraph.add_run(f"{index}. {affil}")

        if affil != list(affiliations_dict.keys())[-1] :
            affliations_paragraph.add_run("\n")

    # Save the document
    doc.save("Authors_Affiliations.docx")

    ##
